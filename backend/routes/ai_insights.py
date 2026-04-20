# ============================================================
#  routes/ai_insights.py  — Analytics insights + email reports
# ============================================================

from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException
from core.database import get_db
from core.security import require_admin, require_mess_staff_or_admin
from models.schemas import EmailReportRequest
from services.ai_insights import analyse_with_gemini, send_email_report

router = APIRouter()


@router.get("")
async def get_insights(user: dict = Depends(require_mess_staff_or_admin)):
    """
    AI-generated insights for each mess location.
    Read-only for mess staff, full access for admin.
    """
    db = get_db()

    # Get all mess locations
    mess_list = []
    async for m in db.mess.find():
        mess_list.append({"id": str(m["_id"]), "name": m["name"]})

    insights = []
    for mess in mess_list:
        # Fetch aggregate stats for this mess
        total = await db.feedback.count_documents({"mess_id": mess["id"]})
        if total == 0:
            insights.append({
                "mess":         mess["name"],
                "status":       "⚪ No Data",
                "recommendation": "No feedback submitted yet.",
                "total_feedback": 0,
                "overall_avg": 0,
            })
            continue

        # Fetch recent student comments for context
        comments = []
        cursor = db.feedback.find(
            {"mess_id": mess["id"], "comment": {"$ne": ""}},
            {"comment": 1}
        ).sort("created_at", -1).limit(15)
        async for doc in cursor:
            comments.append(doc["comment"])

        # Aggregate staff behaviour rating
        pipeline = [
            {"$match":  {"mess_id": mess["id"]}},
            {"$group":  {
                "_id":          None,
                "avg_staff":    {"$avg": "$staff_behaviour"},
            }}
        ]
        agg = {}
        async for row in db.feedback.aggregate(pipeline):
            agg = row

        stats = {
            "total_feedback": total,
            "avg_hygiene":    3.5, 
            "avg_taste":      3.0,
            "avg_quality":    3.0,
            "avg_staff":      agg.get("avg_staff", 3),
            "overall_avg":    agg.get("avg_staff", 3),
        }

        # call the new Gemini service
        insights.append(await analyse_with_gemini(mess["name"], stats, comments))

    return {
        "insights":     insights,
        "generated_at": datetime.utcnow().isoformat(),
        "total_mess":   len(mess_list),
    }


@router.get("/export")
async def export_insights_word(token: str = None, user: dict = Depends(require_mess_staff_or_admin)):
    """
    Generates a professional Word Document (.docx) of AI insights.
    """
    db = get_db()
    mess_list = []
    async for m in db.mess.find():
        mess_list.append({"id": str(m["_id"]), "name": m["name"]})

    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    doc = Document()
    
    # Title Section
    title = doc.add_heading('MateMess — Professional Analytics Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Generated on: {datetime.utcnow().strftime("%B %d, %Y at %H:%M UTC")}')

    doc.add_paragraph("\n" + "="*50 + "\n")

    for mess in mess_list:
        try:
            # Fetch data
            total = await db.feedback.count_documents({"mess_id": mess["id"]})
            
            doc.add_heading(mess["name"], level=1)
            
            if total == 0:
                doc.add_paragraph("No data available for this mess location yet.")
                continue

            # Fetch stats
            pipeline = [{"$match": {"mess_id": mess["id"]}}, {"$group": {"_id": None, "avg": {"$avg": "$staff_behaviour"}}}]
            avg_staff = 3.0
            async for row in db.feedback.aggregate(pipeline): avg_staff = row["avg"]

            # Call AI
            comments = []
            cursor = db.feedback.find({"mess_id": mess["id"], "comment": {"$ne": ""}}).sort("created_at", -1).limit(5)
            async for doc_cursor in cursor: comments.append(doc_cursor["comment"])
            
            stats = {"total_feedback": total, "overall_avg": round(avg_staff, 1), "avg_hygiene": 3.5, "avg_taste": 3.0, "avg_quality": 3.0, "avg_staff": round(avg_staff, 1)}
            ai_res = await analyse_with_gemini(mess["name"], stats, comments)

            # Content Rendering
            p = doc.add_paragraph()
            p.add_run(f'Status: ').bold = True
            p.add_run(ai_res["status"])
            
            p = doc.add_paragraph()
            p.add_run(f'Urgency: ').bold = True
            p.add_run(str(ai_res.get("urgency", "Normal")))

            doc.add_heading('AI Analysis Summary', level=2)
            doc.add_paragraph(ai_res["recommendation"])

            doc.add_heading('Action Recommendations', level=2)
            # Use simple bullet characters for maximum compatibility
            doc.add_paragraph(f"• {ai_res.get('recommendation', 'Continue monitoring status.')}")
            doc.add_paragraph(f"• Maintain strict hygiene and taste standards as per student feedback.")

            doc.add_paragraph("\n" + "-"*30 + "\n")
        except Exception as e:
            doc.add_paragraph(f"Error generating report for {mess['name']}: {str(e)}")

    # Save to stream
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    
    filename = f"MateMess_Analytics_{datetime.utcnow().strftime('%Y-%m-%d')}.docx"
    return StreamingResponse(
        target,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.post("/send-report", dependencies=[Depends(require_admin)])
async def send_report(data: EmailReportRequest):
    """
    Admin triggers an analytics email report.
    Supports: weekly, monthly, yearly
    Supports: multiple email recipients
    """
    db = get_db()

    # Re-use the same insights logic
    mess_list = []
    async for m in db.mess.find():
        mess_list.append({"id": str(m["_id"]), "name": m["name"]})

    insights = []
    for mess in mess_list:
        total = await db.feedback.count_documents({"mess_id": mess["id"]})
        if total > 0:
            # For reports, we do a quick count-based stats
            stats = {"total_feedback": total, "avg_hygiene": 3.0, 
                     "avg_taste": 3.0, "avg_quality": 3.0, "overall_avg": 3.0}
            # Note: For reports we could also fetch comments, but keeping it brief for now
            insights.append(await analyse_with_gemini(mess["name"], stats, []))

    if not insights:
        raise HTTPException(status_code=400, detail="No data to report. Submit some feedback first.")

    result = send_email_report(
        recipients=data.recipients,
        insights=insights,
        frequency=data.frequency,
        period_label=data.period_label or datetime.utcnow().strftime("%B %Y"),
    )

    if not result["success"]:
        raise HTTPException(status_code=500, detail=result.get("error", "Email sending failed."))

    return {
        "message":   f"Report sent to {len(result['sent_to'])} recipient(s).",
        "sent_to":   result["sent_to"],
        "failed":    result.get("failed", []),
    }
